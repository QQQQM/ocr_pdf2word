import os
import re
import shutil
import time
import subprocess
from aip import AipOcr
from docx import Document
# from wand.image import Image
from configparser import ConfigParser

# 读取配置文件config.cfg
config_parser = ConfigParser()
config_parser.read('config.cfg')
config = config_parser['default']
tasks = []

client = AipOcr(config['appId'], config['apiKey'], config['secretKey'])   # 连接百度云的服务器
print("连接服务器成功！")


def get_file_content(file):
    with open(file, 'rb') as fp:
        return fp.read()


def img_to_str(image_path):
    image = get_file_content(image_path)
    result = client.basicGeneral(image)
    if 'words_result' in result:
        return '\n'.join([w['words'] for w in result['words_result']])


def mkdir(path):
    is_exists = os.path.exists(path)
    if not is_exists:
        os.makedirs(path)
        print("\n成功创建图片路径：", path)
    else:
        shutil.rmtree(path)
        os.makedirs(path)
        print("\n存在同名图片路径文件夹，成功删除后创建：", path)
    return True


def save_text_to_word(content1, file_path):
    doc = Document()
    for line in content1.split('\n'):
        paragraph = doc.add_paragraph()
        paragraph.add_run(remove_control_characters(line))
    doc.save(file_path)


def remove_control_characters(content2):
    mpa = dict.fromkeys(range(32))
    return content2.translate(mpa)


'''
def convert_pdf_to_img(pdf_filename, img_filename):
    print("\n正在将：", pdf_filename, "\n转为图片存放在：", img_filename)
    with Image(filename=pdf_filename) as img:
        print('总页数为：', len(img.sequence))
        with img.convert(config['img_type']) as converted:
            converted.save(filename=img_filename)
'''


def convert_pdf_to_img_2(pdf_filename, img_filename, img_pathname):
    print("\n正在将：", pdf_filename, "\n转为图片存放在：", img_filename, "\n请稍等片刻。。。")
    subprocess.check_output(['pdfimages.exe', '-png', pdf_filename, img_filename])
    print("\n全部转换完成！")


def html_to_plain_text (text):
    if flag3 == '1':
        op = input("\n请选择替换方式：\n\t1：匹配仅包含全大写单词和空格的一行（eg：THE MATING MIND ）；\n" +
                   "\t2：匹配包含全大写单词,数字和空格的一行（eg：12 THE MATING MIND）；\n" +
                   "\t3：匹配不包含小写字母的一行，会删除中文！（eg：F&  #IRSTANCHORB&OOKSEDITION , APRIL 2001 ）;\n" +
                   "\t0：手动输入正则表达式。\n")
        if op == '1':
            text = re.sub('^[A-Z|\\s]+$\n', '\n', text, flags=re.M | re.S)
        if op == '2':
            text = re.sub('^[A-Z|\\s|0-9]+$\n', '\n', text, flags=re.M | re.S)
        if op == '3':
            text = re.sub('^[^a-z]+$\n', '\n', text, flags=re.M | re.S)
        if op == '0':
            sou = input("\n请输入正则表达式：\n")
            des = input("\n请输入要替换为的内容：\n")
            text = re.sub(sou, des, text, flags=re.M | re.S)
        # text = re.sub('^[A-Z|\s]+$\n', '\n', text, flags=re.M | re.S)1

    if flag == '1':
        text = re.sub(r"([A-Z][a-zA-Z]+)\s*$\n", r"\1!qm! \n", text, flags=re.M | re.S)
        text = re.sub(r"(\.|\?|\!|。|？|！)$\n", r"\1 \n", text, flags=re.M | re.S)
        text = re.sub(r"(?<!((\.|\?|\!|。|？|！)\s))$\n", "", text, flags=re.M | re.S)
        text = re.sub(r"!qm!", r"", text, flags=re.M | re.S)
    if flag == '2':
        text = re.sub(r"(\.|\?|\!|。|？|！)$\n", r"\1 $\n", text, flags=re.M | re.S)
        text = re.sub(r"(?<!((\.|\?|\!|。|？|！)\s))$\n", "", text, flags=re.M | re.S)
    if flag2 == '1':
        text = re.sub(r'\s([a-zA-Z])\s([a-zA-Z])\s', r'\1\2', text, flags=re.M | re.S)
    if flag3 == '2':
        op = input("\n请选择替换方式：\n\t1：匹配仅包含全大写单词和空格的一行（eg：THE MATING MIND ）；\n" +
                   "\t2：匹配包含全大写单词,数字和空格的一行（eg：12 THE MATING MIND）；\n" +
                   "\t3：匹配不包含小写字母的一行（eg：F&  #IRSTANCHORB&OOKSEDITION , APRIL 2001 ）;\n" +
                   "\t0：手动输入正则表达式。\n")
        if op == '1':
            text = re.sub('^[A-Z|\\s]+$\n', '\n', text, flags=re.M | re.S)
        if op == '2':
            text = re.sub('^[A-Z|\\s|0-9]+$\n', '\n', text, flags=re.M | re.S)
        if op == '3':
            text = re.sub('^[^a-z]+$\n', '\n', text, flags=re.M | re.S)
        if op == '0':
            sou = input("\n请输入正则表达式：\n")
            des = input("\n请输入要替换为的内容：\n")
            text = re.sub(sou, des, text, flags=re.M | re.S)
    return text


op0 = input("请选择操作：\n\t1：识别pdf；\n\t0：识别图片。\n")
if op0 == '0':
    flag = input("\n是否需要处理换行？\n" +
                 "\t1：表示保留结尾符号和大写开头单词结尾换行;\n"
                 "\t2：表示仅保留结尾符号结尾换行;\n\t0：表示不处理。\n")
    flag2 = input("\n是否需要处理空格问题？;\n\t1：表示处理空格问题;\n\t0：表示不处理。\n")
    flag3 = input("\n是否需要正则字符串处理？;\n\t1：表示在换行空格处理前处理;\n\t2：表示在换行空格处理后处理;\n\t0：表示不处理。\n")

    img_file = config['img_folder']
    word_file = config['word_folder'] + '/' + 'picture.docx'
    text_sor = ""
    for img in os.listdir(img_file):
        img_extension_name = os.path.splitext(img)[1]
        img_name = img_file + '/' + img
        print("正在识别:", img_name)
        text_tmp = img_to_str(img_name)
        # print("识别完毕，该段文本为： ", text_tmp)
        text_sor += text_tmp
        time.sleep(0.5)

    print("\n图片全部识别完毕!")
    save_text_to_word(text_sor, word_file)
    text_sor = html_to_plain_text(text_sor)
    print("\n文本处理完毕!")
    save_text_to_word(text_sor, word_file)
    print("\n完成！")
    os.system("pause")
    exit(1)


for file in os.listdir(config['pdf_folder']):
    extension_name = os.path.splitext(file)[1]
    if extension_name != '.pdf':
        continue
    print("\n检测到新的pdf！")

    flag = input("\n是否需要处理换行？\n" +
                 "\t1：表示保留结尾符号和大写开头单词结尾换行;\n"
                 "\t2：表示仅保留结尾符号结尾换行;\n\t0：表示不处理。\n")
    flag2 = input("\n是否需要处理空格问题？;\n\t1：表示处理空格问题;\n\t0：表示不处理。\n")
    flag3 = input("\n是否需要正则字符串处理？;\n\t1：表示在换行空格处理前处理;\n\t2：表示在换行空格处理后处理;\n\t0：表示不处理。\n")

    file_name = os.path.splitext(file)[0]
    pdf_file = config['pdf_folder'] + '/' + file
    img_path = config['img_folder']
    img_file = img_path + '/' + '1.' + config['img_type']
    # img_file = config['img_folder']
    if flag2 == '1':
        file_name = 'S+' + file_name
    if flag == '1':
        file_name = 'L1+' + file_name
    if flag == '2':
        file_name = 'L2+' + file_name
    if flag3 == '1':
        file_name = 'D+' + file_name
    word_file = config['word_folder'] + '/' + file_name + '.docx'
    print("文件名：", file_name,"\n源路径：", pdf_file, "\n目标路径：",word_file, "\n图片路径：", img_path)
    mkdir(img_path)
    convert_pdf_to_img_2(pdf_file, img_file, img_path)
    '''
    op = input("请选择转换机制（默认为1）：\n\t1：pdfimages；\n\t2：wand.\n")

    if op == '1':
        convert_pdf_to_img_2(pdf_file, img_file, img_path)
    else:
        convert_pdf_to_img(pdf_file, img_file)
    '''
    text_sor = ""
    for img in os.listdir(img_path):
        img_extension_name = os.path.splitext(img)[1]
        img_name = img_path + '/' + img
        print("正在识别:", img_name)
        text_tmp = img_to_str(img_name)
        # print("识别完毕，该段文本为： ", text_tmp)

        text_sor += text_tmp
        time.sleep(0.5)
    print("\n图片全部识别完毕!")
    save_text_to_word(text_sor, word_file)
    text_sor = html_to_plain_text(text_sor)
    print("\n文本处理完毕!")
    save_text_to_word(text_sor, word_file)
print("\n完成！")
os.system("pause")
exit(1)
